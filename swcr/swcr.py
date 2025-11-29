# -*- coding: utf-8 -*-
import logging
import codecs
from os.path import abspath, relpath
from os import scandir
from dataclasses import dataclass
import argparse
from typing import List, Tuple

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

DEFAULT_INDIRS = ['.']
DEFAULT_EXTS = ['c', 'h', 'py', 'js', 'java', 'cpp', 'hpp']
DEFAULT_COMMENT_CHARS = ['/*', '*', '*/', '//', '#']

def del_slash(dirs):
    return [dir_[:-1] if dir_[-1] == '/' else dir_ for dir_ in dirs]

class CodeFinder(object):
    def __init__(self, exts=None):
        self.exts = exts if exts else DEFAULT_EXTS

    @staticmethod
    def is_hidden_file(file):
        return file[0] == '.'

    @staticmethod
    def should_be_excluded(file, excludes = None):
        if not excludes:
            return False
        if not isinstance(excludes, list):
            excludes = [excludes]
        return any(file.startswith(exclude) for exclude in excludes)

    def is_code(self, file):
        is_code_file = any(file.endswith(ext) for ext in self.exts)
        return is_code_file

    def find(self, indir, excludes = None):
        files = []
        for entry in scandir(indir):
            entry_name = entry.name
            entry_path = abspath(entry.path)
            if self.is_hidden_file(entry_name) or self.should_be_excluded(entry_path, excludes):
                continue
            if entry.is_file():
                if self.is_code(entry_name):
                    files.append(entry_path)
            else:
                files.extend(self.find(entry_path, excludes=excludes))
        logger.debug('%s directory:%d code files.', indir, len(files))
        return files

class PDFCodeWriter(object):
    def __init__(self, font_name='Courier', font_size=7, max_front_pages=30, max_back_pages=30):
        self.font_name = font_name
        self.font_size = font_size
        self.max_front_pages = max_front_pages
        self.max_back_pages = max_back_pages
        self.line_height = font_size + 1
        self.margin_left = 40
        self.margin_right = 40
        self.margin_top = 60
        self.margin_bottom = 40
        self.page_width = A4[0]
        self.page_height = A4[1]
        self.usable_width = self.page_width - self.margin_left - self.margin_right
        self.usable_height = self.page_height - self.margin_top - self.margin_bottom
        self.lines_per_page = int(self.usable_height / self.line_height)
        
        # 存储所有内容行
        self.all_lines = []
        self.canvas = None
        
        # 注册中文字体
        self.setup_fonts()
        
    def setup_fonts(self):
        """设置字体，支持中文"""
        try:
            # 尝试使用系统中文字体
            import platform
            system = platform.system()
            
            if system == "Darwin":  # macOS
                try:
                    # 尝试使用 PingFang SC
                    pdfmetrics.registerFont(TTFont('Chinese', '/System/Library/Fonts/PingFang.ttc'))
                    self.chinese_font = 'Chinese'
                except:
                    try:
                        # 尝试使用 STHeiti
                        pdfmetrics.registerFont(TTFont('Chinese', '/System/Library/Fonts/STHeiti Medium.ttc'))
                        self.chinese_font = 'Chinese'
                    except:
                        # 如果都失败，使用Helvetica，但中文会显示为方框
                        self.chinese_font = 'Helvetica'
                        print("Warning: Chinese font not available, Chinese characters may not display correctly")
            elif system == "Windows":
                try:
                    # Windows 中文字体
                    pdfmetrics.registerFont(TTFont('Chinese', 'C:/Windows/Fonts/simhei.ttf'))
                    self.chinese_font = 'Chinese'
                except:
                    try:
                        pdfmetrics.registerFont(TTFont('Chinese', 'C:/Windows/Fonts/simsun.ttc'))
                        self.chinese_font = 'Chinese'
                    except:
                        self.chinese_font = 'Helvetica'
                        print("Warning: Chinese font not available, Chinese characters may not display correctly")
            else:
                # Linux等其他系统
                try:
                    pdfmetrics.registerFont(TTFont('Chinese', '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc'))
                    self.chinese_font = 'Chinese'
                except:
                    self.chinese_font = 'Helvetica'
                    print("Warning: Chinese font not available, Chinese characters may not display correctly")
        except Exception as e:
            print(f"Font setup error: {e}")
            self.chinese_font = 'Helvetica'

    def contains_chinese(self, text):
        """检查文本是否包含中文字符"""
        for char in text:
            if '\u4e00' <= char <= '\u9fff':
                return True
        return False
        
    def check_file_encoding(self, file_path):
        """ check file encoding """
        import chardet
        with open(file_path, 'rb') as fd:
            raw_data = fd.read()
            result = chardet.detect(raw_data)
            encode_str = result['encoding']
            confidence = result['confidence']
            logging.info("input_file: %s, encoding: %s, confidence: %f", file_path, encode_str, confidence)
            
            # 如果置信度太低，尝试常见编码
            if confidence < 0.7:
                for encoding in ['utf-8', 'gbk', 'gb2312', 'big5']:
                    try:
                        raw_data.decode(encoding)
                        encode_str = encoding
                        break
                    except:
                        continue
            
            return encode_str

    @staticmethod
    def is_blank_line(line):
        return not bool(line.strip())

    def is_comment_line(self, line, comment_chars):
        return any(line.lstrip().startswith(comment_char) for comment_char in comment_chars)
    
    def wrap_long_line(self, line, max_chars=90):
        """将长行拆分为多行"""
        if len(line) <= max_chars:
            return [line]
        
        wrapped_lines = []
        while len(line) > max_chars:
            wrapped_lines.append(line[:max_chars])
            line = line[max_chars:]
        if line:
            wrapped_lines.append(line)
        return wrapped_lines

    def collect_code_lines(self, files, comment_chars, base_dir=None):
        """收集所有代码行"""
        for file in files:
            encoding = self.check_file_encoding(file)
            print(f"Processing: {file}, encoding: {encoding}")
            
            # 添加文件相对路径注释
            if base_dir:
                try:
                    relative_path = relpath(file, base_dir)
                except ValueError:
                    relative_path = file
            else:
                relative_path = file
            
            # 添加文件路径注释行
            self.all_lines.append(f"# File: {relative_path}")
            
            # 读取文件内容
            try:
                with codecs.open(file, 'r', encoding, errors='replace') as fp:
                    for line in fp:
                        line = line.rstrip()
                        # 处理长行换行
                        wrapped_lines = self.wrap_long_line(line, max_chars=90)
                        self.all_lines.extend(wrapped_lines)
            except Exception as e:
                print(f"Error reading file {file}: {e}")
                self.all_lines.append(f"# Error reading file: {e}")
        
        print(f"Total lines collected: {len(self.all_lines)}")

    def count_effective_lines(self, lines, comment_chars):
        """计算有效行数（非空非注释行）"""
        count = 0
        for line in lines:
            if not self.is_blank_line(line) and not self.is_comment_line(line, comment_chars):
                count += 1
        return count

    def split_lines_for_pages(self, comment_chars):
        """将代码行分组为页面，确保每页至少50行有效代码"""
        if not self.all_lines:
            return [], []
        
        # 计算总的有效行数
        total_effective_lines = self.count_effective_lines(self.all_lines, comment_chars)
        print(f"Total effective lines: {total_effective_lines}")
        
        # 每页需要的有效行数（不计空白行）- 确保最终文档每页显示50行
        # 考虑到页眉占2行，实际内容区域需要更多行才能填满页面
        lines_per_page = 52  # 增加到52行以确保页面填满
        
        # 分页逻辑
        front_pages = []
        back_pages = []
        
        current_page_lines = []
        current_effective_count = 0
        page_count = 0
        
        i = 0
        while i < len(self.all_lines) and page_count < self.max_front_pages:
            line = self.all_lines[i]
            current_page_lines.append(line)
            
            # 检查是否为有效行（非空白行）
            if not self.is_blank_line(line):
                current_effective_count += 1
            
            # 如果有效行数达到每页限制，或者到达文件末尾，完成当前页
            if current_effective_count >= lines_per_page or i == len(self.all_lines) - 1:
                front_pages.append(current_page_lines.copy())
                print(f"Front page {page_count + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
                current_page_lines = []
                current_effective_count = 0
                page_count += 1
            
            i += 1
        
        # 如果前面没有用完所有行，需要处理后面的页面
        if i < len(self.all_lines):
            # 找到最后30页的开始位置
            remaining_lines = self.all_lines[i:]
            remaining_effective = self.count_effective_lines(remaining_lines, comment_chars)
            
            if remaining_effective > self.max_back_pages * lines_per_page:
                # 找到最后30页的开始位置
                target_effective_lines = self.max_back_pages * lines_per_page
                start_pos = len(remaining_lines) - 1
                effective_count = 0
                
                # 从后往前计算
                for j in range(len(remaining_lines) - 1, -1, -1):
                    line = remaining_lines[j]
                    if not self.is_blank_line(line):
                        effective_count += 1
                    
                    if effective_count >= target_effective_lines:
                        start_pos = j
                        break
                
                # 从找到的位置开始分页
                back_lines = remaining_lines[start_pos:]
                current_page_lines = []
                current_effective_count = 0
                back_page_num = 0
                
                for line in back_lines:
                    current_page_lines.append(line)
                    
                    if not self.is_blank_line(line):
                        current_effective_count += 1
                    
                    if current_effective_count >= lines_per_page:
                        back_pages.append(current_page_lines.copy())
                        print(f"Back page {back_page_num + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
                        current_page_lines = []
                        current_effective_count = 0
                        back_page_num += 1
                
                # 添加最后一页（如果有剩余内容）
                if current_page_lines:
                    back_pages.append(current_page_lines)
                    print(f"Back page {back_page_num + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
            else:
                # 剩余行数不足30页，全部作为后页
                current_page_lines = []
                current_effective_count = 0
                back_page_num = 0
                
                for line in remaining_lines:
                    current_page_lines.append(line)
                    
                    if not self.is_blank_line(line):
                        current_effective_count += 1
                    
                    if current_effective_count >= lines_per_page:
                        back_pages.append(current_page_lines.copy())
                        print(f"Back page {back_page_num + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
                        current_page_lines = []
                        current_effective_count = 0
                        back_page_num += 1
                
                if current_page_lines:
                    back_pages.append(current_page_lines)
                    print(f"Back page {back_page_num + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
        
        return front_pages, back_pages

    def create_pdf(self, filename, title, version, front_pages, back_pages):
        """创建PDF文件"""
        self.canvas = canvas.Canvas(filename, pagesize=A4)
        
        # 写入前面的页面
        for page_num, page_lines in enumerate(front_pages, 1):
            self.draw_page(page_lines, page_num, title, version)
            self.canvas.showPage()
        
        # 如果有后面的页面，添加省略页
        if back_pages:
            self.draw_ellipsis_page(len(front_pages) + 1, title, version)
            self.canvas.showPage()
            
            # 写入后面的页面
            for page_num, page_lines in enumerate(back_pages, len(front_pages) + 2):
                self.draw_page(page_lines, page_num, title, version)
                self.canvas.showPage()
        
        self.canvas.save()

    def draw_header(self, page_num, title, version):
        """绘制页眉"""
        header_y = self.page_height - 30
        
        # 左侧装订线留白（额外留出20点）
        header_left_margin = self.margin_left + 20
        
        # 绘制软件名称和版本号（左侧）
        header_text = f"{title} {version}"
        
        # 检查是否包含中文，选择合适的字体
        if self.contains_chinese(header_text):
            self.canvas.setFont(self.chinese_font, 10)
        else:
            self.canvas.setFont(self.font_name, 10)
        
        self.canvas.drawString(header_left_margin, header_y, header_text)
        
        # 绘制页码（右侧）
        page_text = f"{page_num}"
        self.canvas.setFont(self.font_name, 10)
        self.canvas.drawRightString(self.page_width - self.margin_right, header_y, page_text)
        
        # 绘制页眉下划线
        line_y = header_y - 5
        self.canvas.line(header_left_margin, line_y, self.page_width - self.margin_right, line_y)

    def draw_page(self, lines, page_num, title, version):
        """绘制一页内容"""
        # 绘制页眉
        self.draw_header(page_num, title, version)
        
        # 代码内容的起始位置
        y_position = self.page_height - self.margin_top
        
        # 绘制代码行
        for line in lines:
            if y_position < self.margin_bottom:
                break
            
            # 检查是否包含中文，选择合适的字体
            if self.contains_chinese(line):
                self.canvas.setFont(self.chinese_font, self.font_size)
            else:
                self.canvas.setFont(self.font_name, self.font_size)
            
            # 左侧装订线留白
            x_position = self.margin_left + 20
            self.canvas.drawString(x_position, y_position, line)
            y_position -= self.line_height

    def draw_ellipsis_page(self, page_num, title, version):
        """绘制省略页"""
        # 绘制页眉
        self.draw_header(page_num, title, version)
        
        # 绘制省略符号
        y_center = self.page_height / 2
        self.canvas.setFont(self.font_name, 24)
        text = "......"
        text_width = self.canvas.stringWidth(text, self.font_name, 24)
        x_center = (self.page_width - text_width) / 2
        
        self.canvas.drawString(x_center, y_center, text)

class DOCXCodeWriter(object):
    """DOCX代码文档生成器"""
    def __init__(self, max_front_pages=30, max_back_pages=30):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        self.max_front_pages = max_front_pages
        self.max_back_pages = max_back_pages
        self.all_lines = []
    
    @staticmethod
    def is_blank_line(line):
        return not bool(line.strip())
    
    def is_comment_line(self, line, comment_chars):
        return any(line.lstrip().startswith(comment_char) for comment_char in comment_chars)
    
    def wrap_long_line(self, line, max_chars=90):
        """将长行拆分为多行"""
        if len(line) <= max_chars:
            return [line]
        
        wrapped_lines = []
        while len(line) > max_chars:
            wrapped_lines.append(line[:max_chars])
            line = line[max_chars:]
        if line:
            wrapped_lines.append(line)
        return wrapped_lines
    
    def check_file_encoding(self, file_path):
        """检查文件编码"""
        import chardet
        with open(file_path, 'rb') as fd:
            raw_data = fd.read()
            result = chardet.detect(raw_data)
            encode_str = result['encoding']
            confidence = result['confidence']
            logging.info("input_file: %s, encoding: %s, confidence: %f", file_path, encode_str, confidence)
            
            if confidence < 0.7:
                for encoding in ['utf-8', 'gbk', 'gb2312', 'big5']:
                    try:
                        raw_data.decode(encoding)
                        encode_str = encoding
                        break
                    except:
                        continue
            
            return encode_str
    
    def collect_code_lines(self, files, comment_chars, base_dir=None):
        """收集所有代码行"""
        for file in files:
            encoding = self.check_file_encoding(file)
            print(f"Processing: {file}, encoding: {encoding}")
            
            if base_dir:
                try:
                    relative_path = relpath(file, base_dir)
                except ValueError:
                    relative_path = file
            else:
                relative_path = file
            
            self.all_lines.append(f"# File: {relative_path}")
            
            try:
                with codecs.open(file, 'r', encoding, errors='replace') as fp:
                    for line in fp:
                        line = line.rstrip()
                        wrapped_lines = self.wrap_long_line(line, max_chars=90)
                        self.all_lines.extend(wrapped_lines)
            except Exception as e:
                print(f"Error reading file {file}: {e}")
                self.all_lines.append(f"# Error reading file: {e}")
        
        print(f"Total lines collected: {len(self.all_lines)}")
    
    def count_effective_lines(self, lines, comment_chars):
        """计算有效行数（非空行）"""
        count = 0
        for line in lines:
            if not self.is_blank_line(line):
                count += 1
        return count
    
    def split_lines_for_pages(self, comment_chars):
        """将代码行分组为页面，确保每页至少50行有效代码"""
        if not self.all_lines:
            return [], []
        
        total_effective_lines = self.count_effective_lines(self.all_lines, comment_chars)
        print(f"Total effective lines: {total_effective_lines}")
        
        lines_per_page = 52  # 增加到52行以确保页面填满
        
        front_pages = []
        back_pages = []
        
        current_page_lines = []
        current_effective_count = 0
        page_count = 0
        
        i = 0
        while i < len(self.all_lines) and page_count < self.max_front_pages:
            line = self.all_lines[i]
            current_page_lines.append(line)
            
            if not self.is_blank_line(line):
                current_effective_count += 1
            
            if current_effective_count >= lines_per_page or i == len(self.all_lines) - 1:
                front_pages.append(current_page_lines.copy())
                print(f"Front page {page_count + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
                current_page_lines = []
                current_effective_count = 0
                page_count += 1
            
            i += 1
        
        if i < len(self.all_lines):
            remaining_lines = self.all_lines[i:]
            remaining_effective = self.count_effective_lines(remaining_lines, comment_chars)
            
            if remaining_effective > self.max_back_pages * lines_per_page:
                target_effective_lines = self.max_back_pages * lines_per_page
                start_pos = len(remaining_lines) - 1
                effective_count = 0
                
                for j in range(len(remaining_lines) - 1, -1, -1):
                    line = remaining_lines[j]
                    if not self.is_blank_line(line):
                        effective_count += 1
                    
                    if effective_count >= target_effective_lines:
                        start_pos = j
                        break
                
                back_lines = remaining_lines[start_pos:]
                current_page_lines = []
                current_effective_count = 0
                back_page_num = 0
                
                for line in back_lines:
                    current_page_lines.append(line)
                    
                    if not self.is_blank_line(line):
                        current_effective_count += 1
                    
                    if current_effective_count >= lines_per_page:
                        back_pages.append(current_page_lines.copy())
                        print(f"Back page {back_page_num + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
                        current_page_lines = []
                        current_effective_count = 0
                        back_page_num += 1
                
                if current_page_lines:
                    back_pages.append(current_page_lines)
                    print(f"Back page {back_page_num + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
            else:
                current_page_lines = []
                current_effective_count = 0
                back_page_num = 0
                
                for line in remaining_lines:
                    current_page_lines.append(line)
                    
                    if not self.is_blank_line(line):
                        current_effective_count += 1
                    
                    if current_effective_count >= lines_per_page:
                        back_pages.append(current_page_lines.copy())
                        print(f"Back page {back_page_num + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
                        current_page_lines = []
                        current_effective_count = 0
                        back_page_num += 1
                
                if current_page_lines:
                    back_pages.append(current_page_lines)
                    print(f"Back page {back_page_num + 1}: {len(current_page_lines)} total lines, {current_effective_count} effective lines")
        
        return front_pages, back_pages
    
    def create_docx(self, filename, title, version, front_pages, back_pages):
        """创建DOCX文件"""
        doc = Document()
        
        # 设置页面边距和行距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.5)
        
        # 写入前面的页面
        for page_num, page_lines in enumerate(front_pages, 1):
            self.add_page_to_doc(doc, page_lines, page_num, title, version)
            if page_num < len(front_pages):
                doc.add_page_break()
        
        # 如果有后面的页面，添加省略页
        if back_pages:
            doc.add_page_break()
            self.add_ellipsis_page_to_doc(doc, len(front_pages) + 1, title, version)
            
            # 写入后面的页面
            for page_num, page_lines in enumerate(back_pages, len(front_pages) + 2):
                doc.add_page_break()
                self.add_page_to_doc(doc, page_lines, page_num, title, version)
        
        doc.save(filename)
    
    def add_page_to_doc(self, doc, lines, page_num, title, version):
        """添加一页内容到文档"""
        from docx.shared import RGBColor
        from docx.oxml.ns import qn
        
        # 添加页眉
        header = f"{title} {version}"
        p = doc.add_paragraph()
        run = p.add_run(header)
        run.font.size = Pt(9)
        
        # 在同一行添加页码（右对齐）
        run = p.add_run(f"\t\t\t\t\t\t\t\t\t\t{page_num}")
        run.font.size = Pt(9)
        
        # 设置段落格式：紧凑行距
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        
        # 添加分隔线
        p = doc.add_paragraph()
        run = p.add_run("_" * 95)
        run.font.size = Pt(8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        
        # 添加代码行 - 使用紧凑格式
        for line in lines:
            p = doc.add_paragraph()
            # 空行也要添加，但用空格占位
            run = p.add_run(line if line.strip() else ' ')
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            
            # 设置紧凑的段落格式
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            # 设置中文字体
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    def add_ellipsis_page_to_doc(self, doc, page_num, title, version):
        """添加省略页到文档"""
        # 添加页眉
        header = f"{title} {version}                                                                                                    {page_num}"
        p = doc.add_paragraph()
        run = p.add_run(header)
        run.font.size = Pt(10)
        
        # 添加分隔线
        p = doc.add_paragraph()
        run = p.add_run("_" * 100)
        
        # 添加省略符号
        for _ in range(20):
            doc.add_paragraph()
        
        p = doc.add_paragraph()
        run = p.add_run("......")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(24)

@dataclass
class MainParams:
    title: str
    version: str
    indirs: list
    exts: list
    comment_chars: list
    font_name: str
    font_size: int
    max_front_pages: int
    max_back_pages: int
    excludes: list
    outfile: str
    verbose: bool

def main(main_params: MainParams):
    title = main_params.title
    version = main_params.version
    indirs = main_params.indirs
    exts = main_params.exts
    comment_chars = main_params.comment_chars
    font_name = main_params.font_name
    font_size = main_params.font_size
    max_front_pages = main_params.max_front_pages
    max_back_pages = main_params.max_back_pages
    excludes = main_params.excludes
    outfile = main_params.outfile
    verbose = main_params.verbose

    if not indirs:
        indirs = DEFAULT_INDIRS
    if not exts:
        exts = DEFAULT_EXTS
    if not comment_chars:
        comment_chars = DEFAULT_COMMENT_CHARS
    if verbose:
        logging.basicConfig(level=logging.DEBUG)

    # 第零步，把所有的路径都转换为绝对路径
    new_indirs = []
    for indir in indirs:
        new_indirs.append(abspath(indir))
    indirs = new_indirs

    excludes = del_slash(
        [abspath(exclude) for exclude in excludes] if excludes else []
    )

    # 第一步，查找代码文件
    finder = CodeFinder(exts)
    files = []
    for indir in indirs:
        for file in finder.find(indir, excludes = excludes):
            files.append(file)
    
    print(f"Found {len(files)} code files")

    # 确定基础目录用于计算相对路径
    base_dir = indirs[0] if len(indirs) == 1 else None
    
    # 判断输出格式
    is_docx = outfile.lower().endswith('.docx')
    
    if is_docx:
        # 第二步，生成DOCX
        if not DOCX_AVAILABLE:
            print("Error: python-docx is not installed. Install it with: pip install python-docx")
            return 1
        
        writer = DOCXCodeWriter(
            max_front_pages=max_front_pages,
            max_back_pages=max_back_pages
        )
        
        # 收集所有代码行
        writer.collect_code_lines(files, comment_chars, base_dir)
        
        # 分页
        front_pages, back_pages = writer.split_lines_for_pages(comment_chars)
        
        print(f"Front pages: {len(front_pages)}")
        print(f"Back pages: {len(back_pages)}")
        
        # 创建DOCX
        writer.create_docx(outfile, title, version, front_pages, back_pages)
        
        print(f"DOCX created: {outfile}")
    else:
        # 第二步，生成PDF
        writer = PDFCodeWriter(
            font_name=font_name,
            font_size=font_size,
            max_front_pages=max_front_pages,
            max_back_pages=max_back_pages
        )
        
        # 收集所有代码行
        writer.collect_code_lines(files, comment_chars, base_dir)
        
        # 分页
        front_pages, back_pages = writer.split_lines_for_pages(comment_chars)
        
        print(f"Front pages: {len(front_pages)}")
        print(f"Back pages: {len(back_pages)}")
        
        # 创建PDF
        writer.create_pdf(outfile, title, version, front_pages, back_pages)
        
        print(f"PDF created: {outfile}")
    
    return 0

def parse_args():
    parser = argparse.ArgumentParser(description='Generate PDF from source code files.')
    parser.add_argument('--title', type=str, default='软件著作权申请材料', help='Title for the document')
    parser.add_argument('--version', type=str, default='V1.0', help='Version of the software')
    parser.add_argument('--indirs', type=str, nargs='+', default=['.'], help='Input directories')
    parser.add_argument('--exts', type=str, nargs='+', help='File extensions')
    parser.add_argument('--comment_chars', type=str, nargs='+', help='Comment characters')
    parser.add_argument('--font_name', type=str, default='Courier', help='Font name')
    parser.add_argument('--font_size', type=int, default=9, help='Font size')
    parser.add_argument('--max_front_pages', type=int, default=30, help='Maximum front pages')
    parser.add_argument('--max_back_pages', type=int, default=30, help='Maximum back pages')
    parser.add_argument('--excludes', type=str, nargs='+', default=[], help='Exclude directories/files')
    parser.add_argument('--outfile', type=str, default='code.pdf', help='Output PDF file')
    parser.add_argument('--verbose', action='store_true', help='Enable verbose logging')

    args = parser.parse_args()
    return args

def cli_main():
    """命令行入口点函数"""
    args = parse_args()
    params = MainParams(
        title=args.title,
        version=args.version,
        indirs=args.indirs,
        exts=args.exts,
        comment_chars=args.comment_chars,
        font_name=args.font_name,
        font_size=args.font_size,
        max_front_pages=args.max_front_pages,
        max_back_pages=args.max_back_pages,
        excludes=args.excludes,
        outfile=args.outfile,
        verbose=args.verbose
    )
    return main(params)

if __name__ == '__main__':
    cli_main()
