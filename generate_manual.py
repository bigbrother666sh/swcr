
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

def create_manual_docx(markdown_file, output_file):
    document = Document()

    # Set default font to SimHei for Chinese characters
    document.styles["Normal"].font.name = "SimHei"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    # Add header
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "AI首席情报官 wiseflow V4.1"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add page number (right aligned)
    # This is a bit tricky with python-docx, usually done with fields which are not directly exposed.
    # For simplicity, we\'ll note it\'s a requirement and might need manual adjustment or a more complex approach.
    # For now, we\'ll just add a placeholder or rely on Word\'s default page numbering.
    # A more robust solution would involve manipulating XML directly or using a template with fields.

    with open(markdown_file, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    
    in_code_block = False
    current_paragraph_lines = []

    for line in lines:
        if line.startswith("# "): # Main title
            if current_paragraph_lines:
                document.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            document.add_heading(line.replace("# ", "").strip(), level=0)
        elif line.startswith("## "): # Section heading
            if current_paragraph_lines:
                document.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            document.add_heading(line.replace("## ", "").strip(), level=1)
        elif line.startswith("### "): # Sub-section heading
            if current_paragraph_lines:
                document.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            document.add_heading(line.replace("### ", "").strip(), level=2)
        elif line.startswith("```"): # Code block start/end
            if current_paragraph_lines:
                document.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            in_code_block = not in_code_block
            if not in_code_block: # End of code block
                document.add_paragraph("") # Add an empty paragraph after code block
        elif in_code_block:
            document.add_paragraph(line)
        elif line.startswith("*   "): # List item
            if current_paragraph_lines:
                document.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            document.add_paragraph(line.replace("*   ", "- ", 1), style="List Bullet")
        elif line.startswith("**[图片占位："): # Image placeholder
            if current_paragraph_lines:
                document.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            document.add_paragraph(line)
        elif line.strip() == "---":
            if current_paragraph_lines:
                document.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            document.add_paragraph("") # Add a separator or new paragraph
        elif line.strip(): # Regular paragraph line
            current_paragraph_lines.append(line.strip())
        else: # Empty line, end of paragraph
            if current_paragraph_lines:
                document.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []

    # Add any remaining paragraph lines
    if current_paragraph_lines:
        document.add_paragraph(" ".join(current_paragraph_lines))

    document.save(output_file)

if __name__ == "__main__":
    create_manual_docx("wiseflow_user_manual.md", "wiseflow_user_manual.docx")


